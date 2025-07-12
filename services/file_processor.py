"""
FileProcessor Service for SmartHire AI
Handles file upload, validation, and text extraction from PDF/DOCX files using OCR fallback.
"""

import os
import re
import logging
from typing import Tuple, Optional

import PyPDF2
from docx import Document
from pdf2image import convert_from_path
import pytesseract

# ✅ Make sure this points to the correct tesseract path on your machine
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


class FileProcessor:
    def __init__(self):
        self.allowed_extensions = {'pdf', 'docx', 'doc'}
        self.max_file_size = 100 * 1024 * 1024  # 100 MB

    def allowed_file(self, filename: str) -> bool:
        ext = os.path.splitext(filename)[1].lower().lstrip('.')
        return ext in self.allowed_extensions

    def validate_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        if not os.path.exists(file_path):
            return False, "File not found"

        size = os.path.getsize(file_path)
        if size > self.max_file_size:
            return False, "File too large"
        if size == 0:
            return False, "File is empty"

        ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        if ext not in self.allowed_extensions:
            return False, f"Unsupported file type. Allowed types: {', '.join(self.allowed_extensions)}"

        return True, None

    def extract_text(self, file_path: str) -> str:
        logger.info(f"🧾 Starting text extraction for: {file_path}")
        is_valid, error = self.validate_file(file_path)
        if not is_valid:
            logger.warning(f"❌ File validation failed: {error}")
            raise ValueError(error)

        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.pdf':
                text = self._extract_text_from_pdf(file_path)
                if not text.strip():
                    logger.warning("⚠️ PDF might be scanned. Trying OCR fallback...")
                    text = self._extract_text_from_pdf_ocr(file_path)
                return text or ""

            elif ext in ['.docx', '.doc']:
                return self._extract_text_from_docx(file_path) or ""

            else:
                raise ValueError("Unsupported file extension.")

        except Exception as e:
            logger.error(f"❌ extract_text() failed: {e}", exc_info=True)
            return ""

    def _extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    logger.info(f"📄 Page {i + 1}: {len(page_text or '')} characters extracted")
                    if page_text:
                        text += page_text + "\n"

            logger.info(f"✅ PDF extraction successful — Total: {len(text)} characters")
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {e}", exc_info=True)
            return ""

    def _extract_text_from_pdf_ocr(self, file_path: str) -> str:
        try:
            pages = convert_from_path(file_path, dpi=300)
            text = ""
            for i, page in enumerate(pages):
                logger.info(f"🔍 OCR processing page {i + 1} of {len(pages)}")
                ocr_text = pytesseract.image_to_string(page)
                text += ocr_text + "\n"

            logger.info(f"✅ OCR completed — Total: {len(text)} characters")
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"❌ OCR failed: {e}", exc_info=True)
            return ""

    def _extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = ""

            for para in doc.paragraphs:
                text += para.text.strip() + "\n"

            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join(cell.text.strip() for cell in row.cells)
                    text += row_text + "\n"

            logger.info(f"✅ DOCX extraction successful — Total: {len(text)} characters")
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {e}", exc_info=True)
            return ""

    def _clean_text(self, text: str) -> str:
        try:
            cleaned = re.sub(r'\s+', ' ', text)  # Normalize whitespace
            cleaned = cleaned.replace('\x00', '').replace('\ufeff', '')  # Remove hidden chars
            return cleaned.strip()
        except Exception as e:
            logger.error(f"⚠️ Text cleaning failed: {e}", exc_info=True)
            return text or ""
